# admin.py
from flask import Blueprint, jsonify, request, current_app, Response
from flask_login import login_required, current_user
from sqlalchemy import func, desc, or_
from collections import Counter, defaultdict
import re
import json
import io
from datetime import datetime

from models import (
    SessionLocal,
    Conversation,
    Message,
    User,
    Role,
    user_roles,
    Patient,
    ConversationDiseaseLikelihood,
    create_patient,
    get_next_global_patient_identifier,
    delete_conversation_by_id,
)

# Optional: FAISS-driven disease likelihoods
from medical_case_faiss import MedicalCaseFAISS

admin_bp = Blueprint("admin", __name__, url_prefix="/admin")

# --------------------------
# Auth guards
# --------------------------
def _require_admin():
    return current_user.is_authenticated and any(r.name == "admin" for r in current_user.roles)

def admin_guard():
    if not _require_admin():
        return jsonify({"ok": False, "error": "Admin only"}), 403


def _user_display_name(user) -> str:
    """Return display name for admin: username or email prefix (no raw email)."""
    if user is None:
        return "—"
    return (user.username or "").strip() or (
        (user.email.split("@")[0] if user.email else "User")
    )


# --------------------------
# Helpers: text cleaning & symptom extraction
# --------------------------
# For pulling a single target symptom from recommender text (your existing heuristic)
SYM_RE = re.compile(r"(?:symptom|target|focus)\s*:\s*([A-Za-z][\w\s/-]{1,80})", re.IGNORECASE)

def _extract_symptom(text: str) -> str | None:
    if not text:
        return None
    m = SYM_RE.search(text)
    if m:
        return m.group(1).strip()
    for kw in (
        "headache", "chest pain", "cough", "wheezing", "shortness of breath",
        "fever", "nausea", "dizziness", "fatigue", "joint pain"
    ):
        if kw in text.lower():
            return kw
    return None

# Strip legacy HTML if any
TAG_RE = re.compile(r"<[^>]+>")

def _safe_text(m: Message) -> str:
    msg = getattr(m, "message", "") or ""
    return TAG_RE.sub("", msg)

# Counter-based symptom extraction for tallies/graphs
SYMPTOM_LEXICON = [
    "fever","cough","wheezing","shortness of breath","breathlessness","chest pain","headache",
    "nausea","vomiting","fatigue","dizziness","joint pain","swelling","stiffness","back pain",
    "sore throat","runny nose","rash","abdominal pain","diarrhea","constipation","weight loss",
    "night sweats","palpitations","fainting","tingling","numbness","weakness","pain"
]
CANON = {s: s for s in SYMPTOM_LEXICON}
CANON.update({
    "sob": "shortness of breath",
    "dyspnea": "shortness of breath",
    "tiredness": "fatigue",
    "lightheadedness": "dizziness",
    "chest tightness": "chest pain",
    "loose stools": "diarrhea",
    "constipated": "constipation",
    "weightloss": "weight loss",
})

def extract_symptoms(text: str) -> Counter:
    t = " " + (text or "").lower() + " "
    counts = Counter()
    # phrase-first to catch multi-word entries
    for phrase in sorted(CANON.keys(), key=len, reverse=True):
        pattern = r'\b' + re.escape(phrase) + r'\b'
        hits = re.findall(pattern, t)
        if hits:
            counts[CANON[phrase]] += len(hits)
            # remove to avoid double-counting overlaps
            t = re.sub(pattern, " ", t)
    return counts

# Lazy FAISS loader for disease likelihoods
_faiss = None
def get_faiss():
    global _faiss
    if _faiss is None:
        idx_path  = current_app.config.get('FAISS_INDEX_PATH', 'medical_cases.index')
        meta_path = current_app.config.get('FAISS_METADATA_PATH', 'medical_cases_metadata.pkl')
        f = MedicalCaseFAISS()
        f.load_index(idx_path, meta_path)
        _faiss = f
    return _faiss

# --------------------------
# Overview stats
# --------------------------
@admin_bp.get("/api/summary")
@login_required
def summary():
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        total_users = db.query(User).count()
        clinicians = (
            db.query(User)
              .join(user_roles, user_roles.c.user_id == User.id)
              .join(Role, Role.id == user_roles.c.role_id)
              .filter(Role.name == "clinician")
              .count()
        )
        admins = (
            db.query(User)
              .join(user_roles, user_roles.c.user_id == User.id)
              .join(Role, Role.id == user_roles.c.role_id)
              .filter(Role.name == "admin")
              .count()
        )
        total_convos = db.query(Conversation).count()
        total_messages = db.query(Message).count()
        patient_msgs = db.query(Message).filter(Message.role == "patient").count()
        clinician_msgs = db.query(Message).filter(Message.role == "clinician").count()
        rec_questions = db.query(Message).filter(Message.type == "question_recommender").count()

        # Conversations per day (last 30 rows by date asc)
        convs_per_day = (
            db.query(func.date(Conversation.created_at), func.count(Conversation.id))
              .group_by(func.date(Conversation.created_at))
              .order_by(func.date(Conversation.created_at))
              .limit(30)
              .all()
        )

        # Top clinicians by # of conversations (Conversation.owner_user_id, not ConversationOwner)
        top_clinician_rows = (
            db.query(User, func.count(Conversation.id).label("cnt"))
              .join(user_roles, user_roles.c.user_id == User.id)
              .join(Role, Role.id == user_roles.c.role_id)
              .filter(Role.name == "clinician")
              .outerjoin(Conversation, Conversation.owner_user_id == User.id)
              .group_by(User.id, User.email, User.username)
              .order_by(desc(func.count(Conversation.id)))
              .limit(10)
              .all()
        )
        top_clinicians = [
            {"display_name": _user_display_name(u), "count": c}
            for u, c in top_clinician_rows
        ]

        return jsonify({
            "ok": True,
            "users": {"total": total_users, "clinicians": clinicians, "admins": admins},
            "conversations": {"total": total_convos},
            "messages": {
                "total": total_messages,
                "patient": patient_msgs,
                "clinician": clinician_msgs,
                "recommended": rec_questions
            },
            "series": {
                "conversations_per_day": [[str(d), c] for d, c in convs_per_day],
                "top_clinicians": top_clinicians,
            }
        })
    finally:
        db.close()

# --------------------------
# List clinicians (with conversation counts)
# --------------------------
@admin_bp.get("/api/clinicians")
@login_required
def clinicians():
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        # Count conversations by Conversation.owner_user_id (not ConversationOwner table)
        rows = (
            db.query(User, func.count(Conversation.id).label("convos"))
              .join(user_roles, user_roles.c.user_id == User.id)
              .join(Role, Role.id == user_roles.c.role_id)
              .filter(Role.name == "clinician")
              .outerjoin(Conversation, Conversation.owner_user_id == User.id)
              .group_by(User.id, User.email, User.username)
              .order_by(desc("convos"))
              .all()
        )
        return jsonify({"ok": True, "clinicians": [
            {"id": u.id, "display_name": _user_display_name(u), "email": u.email or "", "conversations": c}
            for u, c in rows
        ]})
    finally:
        db.close()

# --------------------------
# Paginated conversations (owner display name, patient label; optional clinician filter)
# --------------------------
@admin_bp.get("/api/conversations")
@login_required
def conversations():
    """Paginated list of all conversations with owner display name and patient label (admin-only)."""
    if not _require_admin():
        return admin_guard()

    page = int(request.args.get("page", 1))
    size = min(int(request.args.get("size", 20)), 100)
    clinician_id = request.args.get("clinician_id", type=int)
    patient_id = request.args.get("patient_id", type=int)
    offset = (page - 1) * size

    db = SessionLocal()
    try:
        q = db.query(Conversation)
        if clinician_id is not None:
            q = q.filter(Conversation.owner_user_id == clinician_id)
        if patient_id is not None:
            q = q.filter(Conversation.patient_id == patient_id)
        total = q.count()

        rows = (
            db.query(
                Conversation.id,
                Conversation.created_at,
                User.email,
                User.username,
                Conversation.owner_user_id,
                Conversation.patient_id,
                func.count(Message.id).label("message_count"),
            )
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Message, Message.conversation_id == Conversation.id)
        )
        if clinician_id is not None:
            rows = rows.filter(Conversation.owner_user_id == clinician_id)
        if patient_id is not None:
            rows = rows.filter(Conversation.patient_id == patient_id)
        rows = (
            rows.group_by(
                Conversation.id,
                Conversation.created_at,
                User.email,
                User.username,
                Conversation.owner_user_id,
                Conversation.patient_id,
            )
            .order_by(Conversation.created_at.desc())
            .offset(offset)
            .limit(size)
            .all()
        )

        # Prefer stable patient identifiers (e.g. P010) over per-clinician "Patient N"
        patient_ids = {pid for (_cid, _created, _email, _username, _owner_id, pid, _mc) in rows if pid}
        patient_meta: dict[int, dict[str, str]] = {}
        if patient_ids:
            p_rows = (
                db.query(Patient.id, Patient.identifier, Patient.display_name)
                  .filter(Patient.id.in_(patient_ids))
                  .all()
            )
            for pid, ident, disp in p_rows:
                ident_s = (ident or "").strip()
                disp_s = (disp or "").strip()
                label = ident_s or disp_s or "Patient"
                if ident_s and disp_s:
                    label = f"{ident_s} ({disp_s})"
                patient_meta[int(pid)] = {"label": label}

        convs = []
        for (cid, created, email, username, owner_id, patient_id, msg_count) in rows:
            display_name = (username or "").strip() or (
                (email.split("@")[0] if email else "User")
            ) if (email or username) else (str(owner_id) if owner_id is not None else "—")
            patient_label = "—"
            if patient_id:
                patient_label = (patient_meta.get(int(patient_id), {}) or {}).get("label") or "Patient"

            convs.append({
                "id": cid,
                "created_at": created.isoformat(),
                "owner_display_name": display_name,
                "owner_user_id": owner_id,
                "patient_id": patient_id,
                "patient_label": patient_label,
                "message_count": msg_count,
            })

        return jsonify({"ok": True, "page": page, "size": size, "total": total, "conversations": convs})
    finally:
        db.close()


@admin_bp.delete("/api/conversation/<cid>")
@login_required
def delete_conversation(cid):
    """Delete a conversation (and its messages) as admin."""
    if not _require_admin():
        return admin_guard()
    ok = delete_conversation_by_id(cid)
    if not ok:
        return jsonify({"ok": False, "error": "Not found"}), 404
    return jsonify({"ok": True})


# --------------------------
# User Management
# --------------------------
@admin_bp.get("/api/users")
@login_required
def list_users():
    """List all users with their roles."""
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        users = db.query(User).order_by(User.id.desc()).all()
        result = []
        for u in users:
            result.append({
                "id": u.id,
                "email": u.email,
                "username": u.username,
                "display_name": _user_display_name(u),
                "roles": [r.name for r in u.roles],
                "created_at": u.created_at.isoformat() if hasattr(u, 'created_at') and u.created_at else None
            })
        return jsonify({"ok": True, "users": result})
    finally:
        db.close()


@admin_bp.post("/api/users")
@login_required
def create_user():
    """Create a new user with roles."""
    if not _require_admin():
        return admin_guard()

    data = request.get_json(force=True, silent=True) or {}
    email = (data.get("email") or "").strip()
    username = (data.get("username") or "").strip()
    password = (data.get("password") or "").strip()
    role_names = data.get("roles") or []

    if not email:
        return jsonify({"ok": False, "error": "Email is required"}), 400
    if not password:
        return jsonify({"ok": False, "error": "Password is required"}), 400

    db = SessionLocal()
    try:
        # Check if user already exists
        existing = db.query(User).filter(User.email == email).first()
        if existing:
            return jsonify({"ok": False, "error": "User with this email already exists"}), 400

        # Create user
        from security import hash_password
        new_user = User(
            email=email,
            username=username or None,
            password_hash=hash_password(password)
        )
        db.add(new_user)
        db.flush()

        # Assign roles
        for role_name in role_names:
            role = db.query(Role).filter(Role.name == role_name).first()
            if role:
                new_user.roles.append(role)

        db.commit()
        return jsonify({
            "ok": True,
            "user_id": new_user.id,
            "email": new_user.email
        })
    except Exception as e:
        db.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        db.close()


@admin_bp.put("/api/users/<int:user_id>")
@login_required
def update_user(user_id):
    """Update user roles."""
    if not _require_admin():
        return admin_guard()

    data = request.get_json(force=True, silent=True) or {}
    role_names = data.get("roles") or []
    username = data.get("username")

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({"ok": False, "error": "User not found"}), 404

        # Update username if provided
        if username is not None:
            user.username = username.strip() or None

        # Update roles
        user.roles = []
        for role_name in role_names:
            role = db.query(Role).filter(Role.name == role_name).first()
            if role:
                user.roles.append(role)

        db.commit()
        return jsonify({"ok": True})
    except Exception as e:
        db.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        db.close()


@admin_bp.delete("/api/users/<int:user_id>")
@login_required
def delete_user(user_id):
    """Delete a user."""
    if not _require_admin():
        return admin_guard()

    # Prevent deleting yourself
    if current_user.id == user_id:
        return jsonify({"ok": False, "error": "Cannot delete your own account"}), 400

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({"ok": False, "error": "User not found"}), 404

        db.delete(user)
        db.commit()
        return jsonify({"ok": True})
    except Exception as e:
        db.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        db.close()


@admin_bp.get("/api/roles")
@login_required
def list_roles():
    """List all available roles."""
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        roles = db.query(Role).all()
        return jsonify({
            "ok": True,
            "roles": [{"id": r.id, "name": r.name} for r in roles]
        })
    finally:
        db.close()


# --------------------------
# Admin: list all patients (with clinician display name)
# --------------------------
@admin_bp.get("/api/patients")
@login_required
def admin_list_patients():
    """List all patients with clinician display name (admin-only). Optional ?clinician_id=N filters by clinician."""
    if not _require_admin():
        return admin_guard()

    clinician_id = request.args.get("clinician_id", type=int)

    db = SessionLocal()
    try:
        q = (
            db.query(Patient, User.email, User.username)
            .outerjoin(User, User.id == Patient.clinician_id)
            .order_by(Patient.id.desc())
        )
        if clinician_id is not None:
            q = q.filter(Patient.clinician_id == clinician_id)
        rows = q.all()
        patients = []
        for p, email, username in rows:
            clin_name = (username or "").strip() or ((email.split("@")[0] if email else "—") if email else "—")
            patients.append({
                "id": p.id,
                "identifier": p.identifier or "—",
                "display_name": p.display_name or "—",
                "clinician_id": p.clinician_id,
                "clinician_display_name": clin_name if (email or username) else "—",
                "created_at": p.created_at.isoformat() if p.created_at else None,
            })
        return jsonify({"ok": True, "patients": patients})
    finally:
        db.close()


# --------------------------
# Admin: create patient (identifier continues from latest in DB)
# --------------------------
@admin_bp.post("/api/patients")
@login_required
def admin_create_patient():
    """Create a patient; assign to a clinician. Identifier is next global P001, P002, ..."""
    if not _require_admin():
        return admin_guard()
    data = request.get_json(force=True, silent=True) or {}
    clinician_id = data.get("clinician_id")
    if clinician_id is None:
        return jsonify({"ok": False, "error": "clinician_id required"}), 400
    try:
        clinician_id = int(clinician_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "error": "clinician_id must be an integer"}), 400
    identifier = (data.get("identifier") or "").strip()
    if not identifier:
        identifier = get_next_global_patient_identifier()
    display_name = (data.get("display_name") or "").strip() or None
    pid = create_patient(identifier=identifier, clinician_id=clinician_id, display_name=display_name)
    return jsonify({"ok": True, "patient_id": pid, "identifier": identifier})


# --------------------------
# Conversation detail (messages + recommended questions)
# --------------------------
@admin_bp.get("/api/conversation/<cid>")
@login_required
def conversation_detail(cid):
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        msgs = (
            db.query(Message)
              .filter(Message.conversation_id == cid)
              .order_by(Message.created_at.asc())
              .all()
        )

        out_msgs, recos = [], []
        for m in msgs:
            text = _safe_text(m)
            out_msgs.append({
                "id": m.id,
                "role": m.role,
                "type": m.type,
                "text": text,
                "timestamp": m.timestamp,
                "created_at": m.created_at.isoformat(),
            })
            if (m.type == "question_recommender") or (m.role == "Question Recommender"):
                recos.append({
                    "id": m.id,
                    "question": text,
                    "symptom": _extract_symptom(text)
                })

        return jsonify({"ok": True, "messages": out_msgs, "recommended_questions": recos})
    finally:
        db.close()

# --------------------------
# Symptom tallies (global + per-conversation)
# --------------------------
@admin_bp.get("/api/symptoms")
@login_required
def symptoms_api():
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        # Pull all conversations + owners + patients in one pass
        convo_rows = (
            db.query(
                Conversation.id,
                Conversation.created_at,
                User.email,
                User.username,
                Conversation.owner_user_id,
                Conversation.patient_id,
                Patient.identifier,
                Patient.display_name,
            )
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Patient, Patient.id == Conversation.patient_id)
            .order_by(Conversation.created_at.desc())
            .all()
        )
        conv_ids = [r[0] for r in convo_rows]

        def _display(e, un, uid):
            return (un or "").strip() or ((e.split("@")[0] if e else "User")) if (e or un) else (str(uid) if uid is not None else "—")

        owner_map = {}
        for (cid, created, email, username, uid, pid, p_ident, p_display) in convo_rows:
            owner_map[cid] = {
                "display_name": _display(email, username, uid),
                "id": uid,
                "created_at": created.isoformat(),
                "patient_id": pid,
                "patient_identifier": (p_ident or "—") if p_ident else "—",
                "patient_display_name": (p_display or "—") if p_display else "—",
            }

        # No conversations yet
        if not conv_ids:
            return jsonify({"ok": True, "global": {}, "by_conversation": []})

        # Only patient utterances for counting (be forgiving on casing)
        from sqlalchemy import or_
        msgs = (
            db.query(Message)
              .filter(Message.conversation_id.in_(conv_ids))
              .filter(or_(Message.role == "patient", Message.role == "Patient"))
              .order_by(Message.created_at.asc())
              .all()
        )

        from collections import Counter, defaultdict
        global_counts = Counter()
        per_conv = defaultdict(Counter)

        for m in msgs:
            counts = extract_symptoms(m.message or "")  # uses your helper defined above
            global_counts.update(counts)
            per_conv[m.conversation_id].update(counts)

        by_conv = []
        for cid in conv_ids:
            meta = owner_map.get(cid, {})
            by_conv.append({
                "conversation_id": cid,
                "owner_display_name": meta.get("display_name", "—"),
                "owner_user_id": meta.get("id"),
                "patient_id": meta.get("patient_id"),
                "patient_identifier": meta.get("patient_identifier", "—"),
                "patient_display_name": meta.get("patient_display_name", "—"),
                "created_at": meta.get("created_at"),
                "symptoms": dict(per_conv[cid].most_common()),
            })

        return jsonify({
            "ok": True,
            "global": dict(global_counts.most_common()),
            "by_conversation": by_conv
        })
    finally:
        db.close()


# --------------------------
# Disease likelihoods per conversation (FAISS-weighted)
# --------------------------
@admin_bp.get("/api/conversation/<cid>/disease_likelihoods")
@login_required
def conversation_disease_likelihoods(cid):
    if not _require_admin():
        return admin_guard()

    db = SessionLocal()
    try:
        force = (request.args.get("force") or "").strip().lower() in ("1", "true", "yes", "y")

        # Fetch conversation with owner and patient for display
        conv = (
            db.query(Conversation, User.email, User.username, Patient.identifier, Patient.display_name)
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Patient, Patient.id == Conversation.patient_id)
            .filter(Conversation.id == cid)
            .first()
        )
        patient_label = "—"
        clinician_label = "—"
        if conv:
            _c, uemail, uname, pident, pdisplay = conv
            clinician_label = _user_display_name(User(email=uemail, username=uname)) if (uemail or uname) else "—"
            p_id = (pident or "").strip() if pident else ""
            p_disp = (pdisplay or "").strip() if pdisplay else ""
            if p_id and p_disp:
                patient_label = f"{p_id} ({p_disp})"
            elif p_id:
                patient_label = p_id
            elif p_disp:
                patient_label = p_disp

        # If we already analyzed this conversation, return persisted snapshot unless forced
        existing = db.query(ConversationDiseaseLikelihood).filter(ConversationDiseaseLikelihood.conversation_id == cid).first()
        if existing is not None and not force:
            try:
                symptoms = json.loads(existing.symptoms_json) if existing.symptoms_json else {}
            except Exception:
                symptoms = {}
            try:
                top_diseases = json.loads(existing.top_diseases_json) if existing.top_diseases_json else []
            except Exception:
                top_diseases = []
            try:
                faiss_matches = json.loads(existing.faiss_matches_json) if existing.faiss_matches_json else []
            except Exception:
                faiss_matches = []

            return jsonify({
                "ok": True,
                "conversation_id": cid,
                "patient_label": existing.patient_label or patient_label or "—",
                "clinician_label": existing.clinician_label or clinician_label or "—",
                "symptoms": symptoms,
                "top_diseases": top_diseases,
                "cancer_likelihood_pct": existing.cancer_likelihood_pct,
                "analyzed_at": existing.analyzed_at.isoformat() if existing.analyzed_at else None,
                "source": "db",
                "faiss_matches": faiss_matches,
            })

        msgs = (
            db.query(Message)
              .filter(Message.conversation_id == cid)
              .order_by(Message.created_at.asc())
              .all()
        )
        if not msgs:
            return jsonify({"ok": False, "error": "No messages for conversation"}), 404

        # Prefer patient text; fall back to full transcript
        patient_text = " ".join((m.message or "") for m in msgs if (m.role or "").lower() == "patient").strip()
        if not patient_text:
            patient_text = " ".join((m.message or "") for m in msgs if m.message).strip()

        f = get_faiss()
        # be lenient to get a spread of candidates
        results = f.search_similar_cases(patient_text, k=8, similarity_threshold=0.05)

        weights = defaultdict(float)
        for r in results:
            sim = max(float(r.similarity_score), 0.0)
            sus = r.Suspected_illness or {}
            if isinstance(sus, dict):
                for disease, _val in sus.items():
                    if (disease or "").strip():
                        weights[disease.strip()] += sim
            elif isinstance(sus, str) and sus.strip():
                weights[sus.strip()] += sim

        total = sum(weights.values()) or 1.0
        ranked = sorted(
            ({"disease": k, "weight": v, "likelihood_pct": round(100.0 * v / total, 1)} for k, v in weights.items()),
            key=lambda x: (-x["weight"], x["disease"])
        )[:5]

        # Cancer likelihood: similarity-weighted score from cases with cancer diseases or cancer red flags
        total_sim = sum(max(float(r.similarity_score), 0.0) for r in results) or 1.0
        cancer_sim = 0.0
        for r in results:
            sim = max(float(r.similarity_score), 0.0)
            is_cancer = False
            sus = r.Suspected_illness or {}
            if isinstance(sus, dict):
                for disease in sus.keys():
                    if disease and "cancer" in str(disease).lower():
                        is_cancer = True
                        break
            elif isinstance(sus, str) and "cancer" in sus.lower():
                is_cancer = True
            rf = getattr(r, "red_flags", None) or {}
            if isinstance(rf, dict) and rf.get("Possible cancer-related bleeding"):
                is_cancer = True
            if is_cancer:
                cancer_sim += sim
        cancer_likelihood_pct = round(100.0 * cancer_sim / total_sim, 1)

        sym = extract_symptoms(patient_text)

        faiss_matches = [{
            "case_id": r.case_id,
            "similarity": round(float(r.similarity_score), 4),
            "suspected": r.Suspected_illness
        } for r in results]

        # Persist snapshot (upsert)
        snap = existing or ConversationDiseaseLikelihood(conversation_id=cid)
        snap.analyzed_at = datetime.utcnow()
        snap.cancer_likelihood_pct = float(cancer_likelihood_pct) if cancer_likelihood_pct is not None else None
        snap.symptoms_json = json.dumps(dict(sym.most_common()))
        snap.top_diseases_json = json.dumps(ranked)
        snap.faiss_matches_json = json.dumps(faiss_matches)
        snap.patient_label = patient_label or "—"
        snap.clinician_label = clinician_label or "—"
        db.add(snap)
        db.commit()

        return jsonify({
            "ok": True,
            "conversation_id": cid,
            "patient_label": patient_label or "—",
            "clinician_label": clinician_label or "—",
            "symptoms": dict(sym.most_common()),
            "top_diseases": ranked,
            "cancer_likelihood_pct": cancer_likelihood_pct,
            "analyzed_at": snap.analyzed_at.isoformat() if snap.analyzed_at else None,
            "source": "computed",
            "faiss_matches": faiss_matches,
        })
    finally:
        db.close()


# --------------------------
# Export Utilities
# --------------------------
def _generate_pdf_report(title: str, headers: list, rows: list, subtitle: str = None) -> bytes:
    """Generate a PDF report with table data."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12, textColor=colors.HexColor('#7bc148'))
    elements.append(Paragraph(title, title_style))

    if subtitle:
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.gray, spaceAfter=20)
        elements.append(Paragraph(subtitle, subtitle_style))

    elements.append(Spacer(1, 12))

    # Table data
    table_data = [headers] + rows

    # Calculate column widths
    num_cols = len(headers)
    available_width = 10 * inch
    col_width = available_width / num_cols

    table = Table(table_data, colWidths=[col_width] * num_cols)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7bc148')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e5e7eb')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
    ]))

    elements.append(table)

    # Footer
    elements.append(Spacer(1, 20))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray)
    elements.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - Early Cancer Diagnosis System", footer_style))

    doc.build(elements)
    return buffer.getvalue()


def _generate_word_report(title: str, headers: list, rows: list, subtitle: str = None) -> bytes:
    """Generate a Word document with table data."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=1)
    title_para.runs[0].font.color.rgb = RGBColor(123, 193, 72)

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.runs[0].font.size = Pt(10)
        subtitle_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = str(header)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = str(value) if value else "—"
            row.cells[col_idx].paragraphs[0].runs[0].font.size = Pt(9)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - Early Cancer Diagnosis System")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_conversation_pdf(conv_data: dict, messages: list, patient_label: str, clinician_label: str) -> bytes:
    """Generate a PDF for conversation/chat history."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, spaceAfter=6, textColor=colors.HexColor('#7bc148'))
    elements.append(Paragraph("Conversation Transcript", title_style))

    # Metadata
    meta_style = ParagraphStyle('Meta', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#6b7280'), spaceAfter=4)
    created = conv_data.get('created_at', '')
    if created:
        try:
            created = datetime.fromisoformat(created.replace('Z', '+00:00')).strftime('%Y-%m-%d %H:%M:%S')
        except:
            pass

    elements.append(Paragraph(f"<b>Patient:</b> {patient_label}", meta_style))
    elements.append(Paragraph(f"<b>Clinician:</b> {clinician_label}", meta_style))
    elements.append(Paragraph(f"<b>Date:</b> {created}", meta_style))
    elements.append(Paragraph(f"<b>Messages:</b> {len(messages)}", meta_style))
    elements.append(Spacer(1, 16))

    # Messages
    role_colors = {
        'patient': '#1e40af',
        'clinician': '#166534',
        'Question Recommender': '#92400e',
        'Listener': '#166534',
    }

    for msg in messages:
        role = msg.get('role', 'Unknown')
        text = msg.get('text') or msg.get('message') or ''
        timestamp = msg.get('timestamp', '')
        color = role_colors.get(role, '#4b5563')

        role_style = ParagraphStyle('Role', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor(color), fontName='Helvetica-Bold')
        msg_style = ParagraphStyle('Message', parent=styles['Normal'], fontSize=10, spaceAfter=12, leading=14)

        elements.append(Paragraph(f"{role} [{timestamp}]", role_style))
        # Clean text for PDF
        clean_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        elements.append(Paragraph(clean_text, msg_style))

    # Footer
    elements.append(Spacer(1, 20))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray)
    elements.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - Early Cancer Diagnosis System", footer_style))

    doc.build(elements)
    return buffer.getvalue()


def _generate_conversation_word(conv_data: dict, messages: list, patient_label: str, clinician_label: str) -> bytes:
    """Generate a Word document for conversation/chat history."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Conversation Transcript", level=1)
    title.runs[0].font.color.rgb = RGBColor(123, 193, 72)

    # Metadata
    created = conv_data.get('created_at', '')
    if created:
        try:
            created = datetime.fromisoformat(created.replace('Z', '+00:00')).strftime('%Y-%m-%d %H:%M:%S')
        except:
            pass

    meta = doc.add_paragraph()
    meta.add_run(f"Patient: ").bold = True
    meta.add_run(f"{patient_label}\n")
    meta.add_run(f"Clinician: ").bold = True
    meta.add_run(f"{clinician_label}\n")
    meta.add_run(f"Date: ").bold = True
    meta.add_run(f"{created}\n")
    meta.add_run(f"Messages: ").bold = True
    meta.add_run(f"{len(messages)}")

    doc.add_paragraph()
    doc.add_heading("Messages", level=2)

    role_colors = {
        'patient': RGBColor(30, 64, 175),
        'clinician': RGBColor(22, 101, 52),
        'Question Recommender': RGBColor(146, 64, 14),
        'Listener': RGBColor(22, 101, 52),
    }

    for msg in messages:
        role = msg.get('role', 'Unknown')
        text = msg.get('text') or msg.get('message') or ''
        timestamp = msg.get('timestamp', '')
        color = role_colors.get(role, RGBColor(75, 85, 99))

        para = doc.add_paragraph()
        role_run = para.add_run(f"{role} [{timestamp}]")
        role_run.bold = True
        role_run.font.size = Pt(10)
        role_run.font.color.rgb = color

        msg_para = doc.add_paragraph(text)
        msg_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - Early Cancer Diagnosis System")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------
# Export: Users List
# --------------------------
@admin_bp.get("/api/export/users/<format>")
@login_required
def export_users(format: str):
    """Export users list to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        users = db.query(User).order_by(User.id.desc()).all()

        headers = ['ID', 'Email', 'Username', 'Roles', 'Created']
        rows = []
        for u in users:
            roles = ', '.join(r.name for r in u.roles) if u.roles else '—'
            created = u.created_at.strftime('%Y-%m-%d %H:%M') if hasattr(u, 'created_at') and u.created_at else '—'
            rows.append([str(u.id), u.email or '—', u.username or '—', roles, created])

        title = "Users Report"
        subtitle = f"Total: {len(users)} users"

        if format == 'pdf':
            content = _generate_pdf_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': 'attachment; filename=users_report.pdf'})
        else:
            content = _generate_word_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': 'attachment; filename=users_report.docx'})
    finally:
        db.close()


# --------------------------
# Export: Clinicians List
# --------------------------
@admin_bp.get("/api/export/clinicians/<format>")
@login_required
def export_clinicians(format: str):
    """Export clinicians list to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        rows_data = (
            db.query(User, func.count(Conversation.id).label("convos"))
              .join(user_roles, user_roles.c.user_id == User.id)
              .join(Role, Role.id == user_roles.c.role_id)
              .filter(Role.name == "clinician")
              .outerjoin(Conversation, Conversation.owner_user_id == User.id)
              .group_by(User.id, User.email, User.username)
              .order_by(desc("convos"))
              .all()
        )

        headers = ['ID', 'Display Name', 'Email', 'Conversations']
        rows = []
        for u, convos in rows_data:
            display_name = _user_display_name(u)
            rows.append([str(u.id), display_name, u.email or '—', str(convos)])

        title = "Clinicians Report"
        subtitle = f"Total: {len(rows_data)} clinicians"

        if format == 'pdf':
            content = _generate_pdf_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': 'attachment; filename=clinicians_report.pdf'})
        else:
            content = _generate_word_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': 'attachment; filename=clinicians_report.docx'})
    finally:
        db.close()


# --------------------------
# Export: Patients List
# --------------------------
@admin_bp.get("/api/export/patients/<format>")
@login_required
def export_patients(format: str):
    """Export patients list to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        patients_data = (
            db.query(Patient, User.email, User.username)
            .outerjoin(User, User.id == Patient.clinician_id)
            .order_by(Patient.id.desc())
            .all()
        )

        headers = ['Identifier', 'Display Name', 'Assigned Clinician', 'Created']
        rows = []
        for p, email, username in patients_data:
            clin_name = (username or "").strip() or ((email.split("@")[0] if email else "—") if email else "—")
            created = p.created_at.strftime('%Y-%m-%d %H:%M') if p.created_at else '—'
            rows.append([p.identifier or '—', p.display_name or '—', clin_name, created])

        title = "Patients Report"
        subtitle = f"Total: {len(patients_data)} patients"

        if format == 'pdf':
            content = _generate_pdf_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': 'attachment; filename=patients_report.pdf'})
        else:
            content = _generate_word_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': 'attachment; filename=patients_report.docx'})
    finally:
        db.close()


# --------------------------
# Export: Conversations List
# --------------------------
@admin_bp.get("/api/export/conversations/<format>")
@login_required
def export_conversations(format: str):
    """Export all conversations list to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        rows_data = (
            db.query(
                Conversation.id,
                Conversation.created_at,
                User.email,
                User.username,
                Conversation.patient_id,
                func.count(Message.id).label("message_count"),
            )
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Message, Message.conversation_id == Conversation.id)
            .group_by(Conversation.id, Conversation.created_at, User.email, User.username, Conversation.patient_id)
            .order_by(Conversation.created_at.desc())
            .all()
        )

        # Get patient info
        patient_ids = {pid for (_, _, _, _, pid, _) in rows_data if pid}
        patient_meta = {}
        if patient_ids:
            p_rows = db.query(Patient.id, Patient.identifier, Patient.display_name).filter(Patient.id.in_(patient_ids)).all()
            for pid, ident, disp in p_rows:
                label = (ident or "").strip() or (disp or "").strip() or "Patient"
                patient_meta[int(pid)] = label

        headers = ['ID', 'Clinician', 'Patient', 'Messages', 'Created']
        rows = []
        for (cid, created, email, username, patient_id, msg_count) in rows_data:
            clinician = (username or "").strip() or ((email.split("@")[0] if email else "—") if email else "—")
            patient = patient_meta.get(int(patient_id), "—") if patient_id else "—"
            created_str = created.strftime('%Y-%m-%d %H:%M') if created else '—'
            # Truncate conversation ID for readability
            short_id = str(cid)[:8] + "..." if len(str(cid)) > 8 else str(cid)
            rows.append([short_id, clinician, patient, str(msg_count), created_str])

        title = "Conversations Report"
        subtitle = f"Total: {len(rows_data)} conversations"

        if format == 'pdf':
            content = _generate_pdf_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': 'attachment; filename=conversations_report.pdf'})
        else:
            content = _generate_word_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': 'attachment; filename=conversations_report.docx'})
    finally:
        db.close()


# --------------------------
# Export: Single Conversation Detail (Chat History)
# --------------------------
@admin_bp.get("/api/export/conversation/<cid>/<format>")
@login_required
def export_conversation_detail(cid: str, format: str):
    """Export a single conversation (chat history) to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        # Get conversation with owner and patient
        conv = (
            db.query(Conversation, User.email, User.username, Patient.identifier, Patient.display_name)
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Patient, Patient.id == Conversation.patient_id)
            .filter(Conversation.id == cid)
            .first()
        )

        if not conv:
            return jsonify({"ok": False, "error": "Conversation not found"}), 404

        conversation, uemail, uname, pident, pdisplay = conv
        clinician_label = _user_display_name(User(email=uemail, username=uname)) if (uemail or uname) else "—"

        p_id = (pident or "").strip()
        p_disp = (pdisplay or "").strip()
        if p_id and p_disp:
            patient_label = f"{p_id} ({p_disp})"
        elif p_id:
            patient_label = p_id
        elif p_disp:
            patient_label = p_disp
        else:
            patient_label = "—"

        # Get messages
        msgs = (
            db.query(Message)
              .filter(Message.conversation_id == cid)
              .order_by(Message.created_at.asc())
              .all()
        )

        messages = []
        for m in msgs:
            messages.append({
                "role": m.role,
                "text": _safe_text(m),
                "timestamp": m.timestamp or "",
            })

        conv_data = {
            "id": cid,
            "created_at": conversation.created_at.isoformat() if conversation.created_at else "",
        }

        if format == 'pdf':
            content = _generate_conversation_pdf(conv_data, messages, patient_label, clinician_label)
            filename = f"conversation_{cid[:8]}.pdf"
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': f'attachment; filename={filename}'})
        else:
            content = _generate_conversation_word(conv_data, messages, patient_label, clinician_label)
            filename = f"conversation_{cid[:8]}.docx"
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': f'attachment; filename={filename}'})
    finally:
        db.close()


# --------------------------
# Export: Analytics/Symptoms Report
# --------------------------
@admin_bp.get("/api/export/analytics/<format>")
@login_required
def export_analytics(format: str):
    """Export symptoms analytics to PDF or Word."""
    if not _require_admin():
        return admin_guard()

    if format not in ('pdf', 'docx'):
        return jsonify({"ok": False, "error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    db = SessionLocal()
    try:
        # Get conversations with symptoms
        convo_rows = (
            db.query(
                Conversation.id,
                Conversation.created_at,
                User.email,
                User.username,
                Patient.identifier,
                Patient.display_name,
            )
            .outerjoin(User, User.id == Conversation.owner_user_id)
            .outerjoin(Patient, Patient.id == Conversation.patient_id)
            .order_by(Conversation.created_at.desc())
            .all()
        )

        conv_ids = [r[0] for r in convo_rows]

        # Get patient messages
        msgs = []
        if conv_ids:
            msgs = (
                db.query(Message)
                  .filter(Message.conversation_id.in_(conv_ids))
                  .filter(or_(Message.role == "patient", Message.role == "Patient"))
                  .all()
            )

        # Extract symptoms per conversation
        per_conv = defaultdict(Counter)
        for m in msgs:
            counts = extract_symptoms(m.message or "")
            per_conv[m.conversation_id].update(counts)

        headers = ['Patient', 'Clinician', 'Top Symptoms', 'Created']
        rows = []
        for (cid, created, email, username, p_ident, p_display) in convo_rows:
            clinician = (username or "").strip() or ((email.split("@")[0] if email else "—") if email else "—")
            patient = (p_ident or "").strip() or (p_display or "").strip() or "—"
            symptoms = per_conv.get(cid, Counter())
            top_symptoms = ', '.join(f"{s} ({c})" for s, c in symptoms.most_common(3)) or "None"
            created_str = created.strftime('%Y-%m-%d') if created else '—'
            rows.append([patient, clinician, top_symptoms, created_str])

        title = "Symptoms Analytics Report"
        subtitle = f"Analysis of {len(convo_rows)} conversations"

        if format == 'pdf':
            content = _generate_pdf_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/pdf',
                          headers={'Content-Disposition': 'attachment; filename=analytics_report.pdf'})
        else:
            content = _generate_word_report(title, headers, rows, subtitle)
            return Response(content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          headers={'Content-Disposition': 'attachment; filename=analytics_report.docx'})
    finally:
        db.close()
